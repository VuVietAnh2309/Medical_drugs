import os
from docx import Document
import json


def load_data_from_folder(folder_path):
    documents = []

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        
        # Xử lý tệp .txt
        if filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                documents.append(text)
        
        # Xử lý tệp .docx
        elif filename.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            documents.append(text)

        elif filename.endswith('.json'):
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                # Lấy dữ liệu từ trường "context" nếu tồn tại
                for item in data:
                    documents.append(item['context'])

    return documents

# import os
# import json
# from docx import Document

# def load_data_from_folder(folder_path):
#     documents = []
    
#     # Xử lý các tệp .json trước
#     for filename in os.listdir(folder_path):
#         if filename.endswith('.json'):
#             file_path = os.path.join(folder_path, filename)
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 data = json.load(file)
#                 # Lấy dữ liệu từ trường "context" nếu tồn tại
#                 if 'context' in data:
#                     documents.append(data['context'])
    
#     # Sau đó xử lý các tệp .txt và .docx
#     for filename in os.listdir(folder_path):
#         file_path = os.path.join(folder_path, filename)
        
#         # Xử lý tệp .txt
#         if filename.endswith('.txt'):
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 text = file.read()
#                 documents.append(text)
        
#         # Xử lý tệp .docx
#         elif filename.endswith('.docx'):
#             doc = Document(file_path)
#             text = '\n'.join([para.text for para in doc.paragraphs])
#             documents.append(text)

#     return documents

